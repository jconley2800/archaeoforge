from __future__ import annotations

import csv
import json
import mimetypes
from pathlib import Path
from typing import Any

import pymupdf as fitz
import yaml
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from .db import connect, get_source, replace_pages, upsert_source
from .models import SourceRecord, SourceType
from .project import ProjectPaths
from .util import sha256_file

SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".md", ".csv", ".json", ".geojson", ".yaml", ".yml", ".docx",
    ".xlsx", ".xlsm", ".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff",
    ".obj", ".ply", ".glb", ".gltf",
}


def _load_sidecar(path: Path) -> dict[str, Any]:
    candidates = [
        path.with_name(path.name + ".source.yaml"), path.with_suffix(".source.yaml"),
        path.with_name(path.name + ".source.yml"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return yaml.safe_load(candidate.read_text(encoding="utf-8")) or {}
    return {}


def _infer_source_type(path: Path) -> SourceType:
    suffix = path.suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        return SourceType.photograph
    if suffix in {".obj", ".ply", ".glb", ".gltf"}:
        return SourceType.survey_scan
    if suffix == ".geojson":
        return SourceType.map
    if suffix in {".xlsx", ".xlsm", ".csv"}:
        return SourceType.dataset
    return SourceType.other


def _extract_pdf(path: Path, cache_dir: Path, source_id: str, render_visual_pages: bool) -> list[dict[str, Any]]:
    pages: list[dict[str, Any]] = []
    page_cache = cache_dir / "pages" / source_id
    page_cache.mkdir(parents=True, exist_ok=True)
    with fitz.open(path) as document:
        for index, page in enumerate(document):
            page_number = index + 1
            text = page.get_text("text") or ""
            image_count = len(page.get_images(full=True))
            rendered_path: str | None = None
            if render_visual_pages and (image_count > 0 or len(text.strip()) < 120):
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                output = page_cache / f"page-{page_number:04d}.png"
                pixmap.save(output)
                rendered_path = str(output)
            pages.append(
                {
                    "page_number": page_number, "text_content": text, "char_count": len(text),
                    "image_count": image_count, "rendered_page_path": rendered_path,
                }
            )
    return pages


def _extract_docx(path: Path) -> list[dict[str, Any]]:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(blocks)
    return [{"page_number": 1, "text_content": text, "char_count": len(text), "image_count": 0}]


def _extract_workbook(path: Path) -> list[dict[str, Any]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    pages: list[dict[str, Any]] = []
    try:
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines = [f"SHEET: {sheet.title}"]
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    lines.append("\t".join(values))
            text = "\n".join(lines)
            pages.append({"page_number": index, "text_content": text, "char_count": len(text), "image_count": 0})
    finally:
        workbook.close()
    return pages


def _extract_csv(path: Path) -> list[dict[str, Any]]:
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        text = "\n".join("\t".join(row) for row in csv.reader(handle))
    return [{"page_number": 1, "text_content": text, "char_count": len(text), "image_count": 0}]


def _extract_text(path: Path) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [{"page_number": 1, "text_content": text, "char_count": len(text), "image_count": 0}]


def _extract_structured(path: Path) -> list[dict[str, Any]]:
    parsed = (
        yaml.safe_load(path.read_text(encoding="utf-8", errors="replace"))
        if path.suffix.lower() in {".yaml", ".yml"}
        else json.loads(path.read_text(encoding="utf-8", errors="replace"))
    )
    text = json.dumps(parsed, ensure_ascii=False, indent=2)
    return [{"page_number": 1, "text_content": text, "char_count": len(text), "image_count": 0}]


def _extract_image(path: Path) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    with Image.open(path) as image:
        metadata: dict[str, Any] = {
            "width": image.width, "height": image.height, "mode": image.mode, "format": image.format,
        }
        exif = image.getexif()
        if exif:
            metadata["exif"] = {str(key): str(value) for key, value in exif.items()}
    return ([{"page_number": 1, "text_content": "", "char_count": 0, "image_count": 1}], metadata)


def _extract_pages(
    path: Path, cache_dir: Path, source_id: str, render_visual_pages: bool
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path, cache_dir, source_id, render_visual_pages), {}
    if suffix == ".docx":
        return _extract_docx(path), {}
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_workbook(path), {}
    if suffix == ".csv":
        return _extract_csv(path), {}
    if suffix in {".txt", ".md"}:
        return _extract_text(path), {}
    if suffix in {".json", ".geojson", ".yaml", ".yml"}:
        return _extract_structured(path), {}
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        return _extract_image(path)
    return [], {}


def discover_source_files(project: ProjectPaths) -> list[Path]:
    files: list[Path] = []
    for path in project.sources_dir.rglob("*"):
        if not path.is_file() or ".source." in path.name or path.name.lower() == "readme.md":
            continue
        if path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)
    return sorted(files)


def ingest_project(project: ProjectPaths, *, render_visual_pages: bool = False) -> dict[str, Any]:
    connection = connect(project)
    result: dict[str, Any] = {"discovered": 0, "ingested": 0, "pages": 0, "failed": 0, "errors": []}
    try:
        for path in discover_source_files(project):
            result["discovered"] += 1
            try:
                checksum = sha256_file(path)
                generated_id = f"SRC-{checksum[:12].upper()}"
                sidecar = _load_sidecar(path)
                relative_path = str(path.relative_to(project.root))
                existing_by_path_row = connection.execute(
                    "SELECT * FROM sources WHERE relative_path = ?", (relative_path,)
                ).fetchone()
                existing_by_path = dict(existing_by_path_row) if existing_by_path_row else {}
                source_id = sidecar.get("id") or existing_by_path.get("id") or generated_id
                existing = get_source(connection, source_id) or existing_by_path
                pages, extracted_metadata = _extract_pages(
                    path, project.cache_dir, source_id, render_visual_pages
                )
                source = SourceRecord(
                    id=source_id,
                    relative_path=relative_path,
                    title=sidecar.get("title") or existing.get("title") or path.stem.replace("_", " "),
                    authors=sidecar.get("authors") or existing.get("authors", ""),
                    publication_year=sidecar.get("publication_year") or existing.get("publication_year"),
                    source_type=sidecar.get("source_type") or existing.get("source_type") or _infer_source_type(path),
                    url=sidecar.get("url") or existing.get("url", ""),
                    license=sidecar.get("license") or existing.get("license", ""),
                    sha256=checksum,
                    size_bytes=path.stat().st_size,
                    mime_type=mimetypes.guess_type(path.name)[0] or "application/octet-stream",
                    local_copy=True,
                    notes=sidecar.get("notes") or existing.get("notes", ""),
                    metadata={**extracted_metadata, **(sidecar.get("metadata") or {})},
                )
                upsert_source(connection, source)
                replace_pages(connection, source.id, pages)
                result["ingested"] += 1
                result["pages"] += len(pages)
            except Exception as exc:
                result["failed"] += 1
                result["errors"].append(
                    f"{path.relative_to(project.root)}: {type(exc).__name__}: {exc}"
                )
    finally:
        connection.close()
    return result
